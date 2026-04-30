import streamlit as st
import requests
from datetime import datetime
import socket
import pandas as pd
import xml.etree.ElementTree as ET
from io import BytesIO
import plotly.express as px
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter, A4
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

SERVIDOR = "http://127.0.0.1:5000"
CONTACTO_WHATS = "6331124596"
st.set_page_config(page_title="ContaXpert Pro", layout="wide", page_icon="📊")

# ================= LOGO SVG =================
LOGO_SVG = """
<svg width="200" height="60" viewBox="0 0 200 60" xmlns="http://www.w3.org/2000/svg">
  <defs>
    <linearGradient id="grad1" x1="0%" y1="0%" x2="100%" y2="0%">
      <stop offset="0%" style="stop-color:#667eea;stop-opacity:1" />
      <stop offset="100%" style="stop-color:#764ba2;stop-opacity:1" />
    </linearGradient>
  </defs>
  <rect x="5" y="10" width="40" height="40" rx="8" fill="url(#grad1)"/>
  <path d="M15 25 L25 35 L35 20" stroke="white" stroke-width="3" fill="none" stroke-linecap="round"/>
  <text x="55" y="28" font-family="Arial, sans-serif" font-size="18" font-weight="bold" fill="#667eea">ContaXpert</text>
  <text x="55" y="44" font-family="Arial, sans-serif" font-size="12" fill="#764ba2">PRO</text>
</svg>
"""

# ================= CAMPOS CFDI DISPONIBLES =================
CAMPOS_CFDI = {
    "Comprobante": ["UUID", "Serie", "Folio", "Fecha", "FechaTimbrado", "Version", "TipoDeComprobante", "MetodoPago", "FormaPago", "Moneda", "TipoCambio", "SubTotal", "Descuento", "Total", "LugarExpedicion"],
    "Emisor": ["RFC_Emisor", "Nombre_Emisor", "RegimenFiscal_Emisor"],
    "Receptor": ["RFC_Receptor", "Nombre_Receptor", "DomicilioFiscal_Receptor", "RegimenFiscal_Receptor", "UsoCFDI"],
    "Conceptos": ["ClaveProdServ", "NoIdentificacion", "Cantidad", "ClaveUnidad", "Unidad", "Descripcion", "ValorUnitario", "Importe", "Descuento_Concepto", "ObjetoImp"],
    "Impuestos": ["TotalImpuestosTrasladados", "TotalImpuestosRetenidos", "IVA_Trasladado", "IEPS_Trasladado", "ISR_Retenido", "IVA_Retenido", "IEPS_Retenido"],
    "Complementos": ["Certificado", "NoCertificado", "Sello", "SelloSAT", "RfcProvCertif"]
}

def hay_internet():
    try:
        socket.create_connection(("8.8.8.8", 53), timeout=3)
        return True
    except OSError:
        return False

def procesar_xmls_detalle(archivos_xml, campos_seleccionados):
    ns = {'cfdi': 'http://www.sat.gob.mx/cfd/4','cfdi33': 'http://www.sat.gob.mx/cfd/3','tfd': 'http://www.sat.gob.mx/TimbreFiscalDigital'}
    data = []
    for archivo in archivos_xml:
        try:
            tree = ET.parse(archivo)
            root = tree.getroot()
            version = root.get('Version')
            ns_cfdi = ns['cfdi'] if version == '4.0' else ns['cfdi33']
            fila_base = {'Archivo': archivo.name}

            tfd = root.find(f'.//{{{ns["tfd"]}}}TimbreFiscalDigital')
            if tfd is not None:
                fila_base['UUID'] = tfd.get('UUID', '')
                fila_base['FechaTimbrado'] = tfd.get('FechaTimbrado', '')
                fila_base['SelloSAT'] = tfd.get('SelloSAT', '')
                fila_base['NoCertificado'] = tfd.get('NoCertificadoSAT', '')
                fila_base['RfcProvCertif'] = tfd.get('RfcProvCertif', '')

            for campo in ['Serie','Folio','Fecha','Version','TipoDeComprobante','MetodoPago','FormaPago','Moneda','TipoCambio','SubTotal','Descuento','Total','LugarExpedicion','Certificado','NoCertificado','Sello']:
                if campo in campos_seleccionados:
                    val = root.get(campo, '')
                    fila_base[campo] = float(val) if campo in ['SubTotal','Descuento','Total','TipoCambio'] and val else val

            emisor = root.find(f'.//{{{ns_cfdi}}}Emisor')
            if emisor is not None:
                for campo, attr in [('RFC_Emisor','Rfc'),('Nombre_Emisor','Nombre'),('RegimenFiscal_Emisor','RegimenFiscal')]:
                    if campo in campos_seleccionados:
                        fila_base[campo] = emisor.get(attr, '')

            receptor = root.find(f'.//{{{ns_cfdi}}}Receptor')
            if receptor is not None:
                for campo, attr in [('RFC_Receptor','Rfc'),('Nombre_Receptor','Nombre'),('DomicilioFiscal_Receptor','DomicilioFiscalReceptor'),('RegimenFiscal_Receptor','RegimenFiscalReceptor'),('UsoCFDI','UsoCFDI')]:
                    if campo in campos_seleccionados:
                        fila_base[campo] = receptor.get(attr, '')

            impuestos = root.find(f'.//{{{ns_cfdi}}}Impuestos')
            if impuestos is not None:
                if 'TotalImpuestosTrasladados' in campos_seleccionados:
                    fila_base['TotalImpuestosTrasladados'] = float(impuestos.get('TotalImpuestosTrasladados', 0))
                if 'TotalImpuestosRetenidos' in campos_seleccionados:
                    fila_base['TotalImpuestosRetenidos'] = float(impuestos.get('TotalImpuestosRetenidos', 0))

            conceptos = root.findall(f'.//{{{ns_cfdi}}}Concepto')
            if not conceptos:
                data.append(fila_base)
            else:
                for concepto in conceptos:
                    fila = fila_base.copy()
                    for campo, attr in [('ClaveProdServ','ClaveProdServ'),('NoIdentificacion','NoIdentificacion'),('Cantidad','Cantidad'),('ClaveUnidad','ClaveUnidad'),('Unidad','Unidad'),('Descripcion','Descripcion'),('ValorUnitario','ValorUnitario'),('Importe','Importe'),('Descuento_Concepto','Descuento'),('ObjetoImp','ObjetoImp')]:
                        if campo in campos_seleccionados:
                            val = concepto.get(attr, '')
                            fila[campo] = float(val) if attr in ['Cantidad','ValorUnitario','Importe','Descuento'] and val else val

                    traslados = concepto.findall(f'.//{{{ns_cfdi}}}Traslado')
                    retenciones = concepto.findall(f'.//{{{ns_cfdi}}}Retencion')
                    fila['IVA_Trasladado'] = sum(float(t.get('Importe',0)) for t in traslados if t.get('Impuesto')=='002')
                    fila['IEPS_Trasladado'] = sum(float(t.get('Importe',0)) for t in traslados if t.get('Impuesto')=='003')
                    fila['ISR_Retenido'] = sum(float(r.get('Importe',0)) for r in retenciones if r.get('Impuesto')=='001')
                    fila['IVA_Retenido'] = sum(float(r.get('Importe',0)) for r in retenciones if r.get('Impuesto')=='002')
                    fila['IEPS_Retenido'] = sum(float(r.get('Importe',0)) for r in retenciones if r.get('Impuesto')=='003')
                    data.append(fila)
        except Exception as e:
            data.append({'Archivo': archivo.name, 'Error': str(e)})

    df = pd.DataFrame(data)
    cols_finales = ['Archivo'] + [c for c in campos_seleccionados if c in df.columns]
    return df[cols_finales]

def generar_excel(df):
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='Detalle_CFDI', index=False)
    output.seek(0)
    return output

def generar_pdf_tabla(df):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    elements = []
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=18, textColor=colors.HexColor('#667eea'), spaceAfter=30, alignment=1)
    elements.append(Paragraph("Reporte ContaXpert Pro - CFDI", title_style))
    elements.append(Spacer(1, 12))

    data = [df.columns.tolist()] + df.values.tolist()
    table = Table(data, repeatRows=1)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#667eea')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 9),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('FONTSIZE', (0, 1), (-1, -1), 7),
    ]))
    elements.append(table)
    doc.build(elements)
    buffer.seek(0)
    return buffer

def generar_word_tabla(df):
    buffer = BytesIO()
    doc = Document()
    doc.add_heading('Reporte ContaXpert Pro - CFDI', 0)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            row_cells[i].text = str(row[col])
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ================= ESTILOS =================
st.markdown("""
<style>
.plan-card {padding: 20px; border-radius: 12px; border: 2px solid #e0; background: white; text-align: center; height: 100%; transition: all 0.3s;}
.plan-card:hover {border-color: #667eea; box-shadow: 0 8px 20px rgba(102,126,234,0.3); transform: translateY(-5px);}
.plan-popular {border-color: #667eea; position: relative;}
.badge-popular {position: absolute; top: -12px; right: 20px; background: linear-gradient(135deg,#667eea 0%,#764ba2 100%); color: white; padding: 4px 12px; border-radius: 20px; font-size: 12px; font-weight: bold;}
.precio-grande {font-size: 36px; font-weight: bold; color: #667eea; margin: 10px 0;}
.feature-check {color: #28a745; margin-right: 8px;}
.demo-overlay {background: #fff3cd; border: 2px dashed #ffc107; padding: 15px; border-radius: 8px; text-align: center; margin-bottom: 20px;}
.header-logo {text-align: center; padding: 20px 0;}
</style>
""", unsafe_allow_html=True)

# Sesión
if 'usuario' not in st.session_state:
    st.session_state.usuario = None
    st.session_state.datos = None
if 'comprar' not in st.session_state:
    st.session_state.comprar = False
if 'modo_demo' not in st.session_state:
    st.session_state.modo_demo = False
if 'internet' not in st.session_state:
    st.session_state.internet = hay_internet()
if 'df_convertido' not in st.session_state:
    st.session_state.df_convertido = None
if 'archivos_temp' not in st.session_state:
    st.session_state.archivos_temp = []
if 'campos_seleccionados' not in st.session_state:
    st.session_state.campos_seleccionados = ['UUID','Serie','Folio','Fecha','RFC_Emisor','Nombre_Emisor','RFC_Receptor','Nombre_Receptor','Total','Descripcion','Importe']
if 'descarga_contabilizada' not in st.session_state:
    st.session_state.descarga_contabilizada = False

# ================= MODO DEMO =================
if st.session_state.modo_demo:
    datos_demo = {'nombre': 'Usuario Demo','plan': 'DEMO','ilimitado': False,'disponibles': 0,'usados': 0,'limite': 300,'extra': 0,'vence': 'Modo Demo','id_empresa': 'DEMO'}
    st.sidebar.markdown(f'<div class="header-logo">{LOGO_SVG}</div>', unsafe_allow_html=True)
    st.sidebar.success(f"👤 {datos_demo['nombre']}")
    st.sidebar.caption(f"Usuario: `DEMO1234`")
    st.sidebar.divider()
    st.sidebar.metric("Plan", "DEMO")
    st.sidebar.metric("Disponibles", "0 XML")
    st.sidebar.warning("🎭 Modo Demostración")
    if st.sidebar.button("🚪 Salir de Demo", use_container_width=True):
        st.session_state.modo_demo = False
        st.rerun()
    st.markdown("""<div class="demo-overlay"><h3>🎭 MODO DEMOSTRACIÓN</h3><p>Estás viendo la interfaz completa de ContaXpert Pro. Para convertir archivos necesitas comprar una licencia.</p></div>""", unsafe_allow_html=True)
    st.markdown(f'<div class="header-logo">{LOGO_SVG}</div>', unsafe_allow_html=True)
    st.title(f"Bienvenido {datos_demo['nombre']} 👋")
    st.stop()

# ================= LOGIN / COMPRA =================
if not st.session_state.usuario:
    st.markdown(f'<div class="header-logo">{LOGO_SVG}</div>', unsafe_allow_html=True)
    st.title("🚀 ContaXpert Pro - Automatiza tus XML")
    st.caption("La plataforma profesional para contadores y empresas que procesan facturas CFDI")

    tab1, tab2, tab3 = st.tabs(["🔑 Iniciar Sesión", "💳 Ver Planes", "🏢 Quiénes Somos"])

    with tab1:
        if not st.session_state.internet:
            st.error("⚠️ Sin conexión a internet. Necesitas internet para usar ContaXpert Pro o activar licencias.")
            st.stop()
        col1, col2, col3 = st.columns([1,2,1])
        with col2:
            if st.button("👀 Probar Demo", use_container_width=True):
                st.session_state.modo_demo = True
                st.rerun()
            st.markdown("---")
            usuario = st.text_input("Usuario", placeholder="Ej: JL1234")
            if st.button("Iniciar Sesión", use_container_width=True, type="primary"):
                if not st.session_state.internet:
                    st.error("❌ Sin internet. No puedes iniciar sesión sin conexión al servidor")
                    st.stop()
                try:
                    with st.spinner('Conectando al servidor...'):
                        r = requests.post(f"{SERVIDOR}/api/login", json={"usuario": usuario}, timeout=15)
                        r.raise_for_status()
                        res = r.json()
                    if res['status'] == 'ok':
                        st.session_state.usuario = usuario
                        st.session_state.datos = res
                        st.success("✅ Conectado al servidor")
                        st.rerun()
                    elif res['status'] == 'pendiente_pago':
                        st.warning(f"⏳ {res['msg']}")
                        st.info("Revisa tu correo. Tu cuenta se activará cuando confirmemos tu pago.")
                    else:
                        st.error(f"❌ {res['msg']}")
                except Exception as e:
                    st.error(f"❌ Error: {str(e)}")
            st.markdown("---")
            with st.expander("❓ Recuperar Usuario"):
                email_rec = st.text_input("Email de registro")
                if st.button("Enviar mis usuarios", use_container_width=True):
                    if email_rec:
                        try:
                            r = requests.post(f"{SERVIDOR}/api/recuperar_usuario", json={"email": email_rec}, timeout=10)
                            r.raise_for_status()
                            if r.json()['status'] == 'ok':
                                st.success(f"✅ {r.json()['msg']}")
                            else:
                                st.error(f"❌ {r.json()['msg']}")
                        except Exception as e:
                            st.error(f"❌ Error: {str(e)}")

    with tab2:
        st.header("📦 Planes diseñados para tu volumen de trabajo")
        col1, col2, col3 = st.columns(3)
        with col1:
            st.markdown("""<div class="plan-card"><h3>🥉 BÁSICO</h3><p style="color:#666;">Para pequeños despachos</p><div class="precio-grande">$149</div><p>al mes / $1,999 pago único</p><hr><p style="text-align:left;"><span class="feature-check">✓</span> 300 XML/mes<br><span class="feature-check">✓</span> 1 usuario<br><span class="feature-check">✓</span> Excel con fórmulas<br><span class="feature-check">✓</span> Soporte por email</p></div>""", unsafe_allow_html=True)
        with col2:
            st.markdown("""<div class="plan-card plan-popular"><div class="badge-popular">MÁS POPULAR</div><h3>🥈 PRO</h3><p style="color:#666;">Para despachos medianos</p><div class="precio-grande">$299</div><p>al mes / $3,499 pago único</p><hr><p style="text-align:left;"><span class="feature-check">✓</span> 500 XML/mes<br><span class="feature-check">✓</span> 1 usuario<br><span class="feature-check">✓</span> Excel + PDF + Word<br><span class="feature-check">✓</span> Soporte WhatsApp<br><span class="feature-check">✓</span> Histórico 12 meses</p></div>""", unsafe_allow_html=True)
        with col3:
            st.markdown("""<div class="plan-card"><h3>🥇 VIP</h3><p style="color:#666;">Para corporativos</p><div class="precio-grande">$399</div><p>al mes / $4,999 pago único</p><hr><p style="text-align:left;"><span class="feature-check">✓</span> XML ilimitados<br><span class="feature-check">✓</span> 1 usuario<br><span class="feature-check">✓</span> API de integración<br><span class="feature-check">✓</span> Soporte prioritario 24/7<br><span class="feature-check">✓</span> Respaldos automáticos</p></div>""", unsafe_allow_html=True)
        st.markdown("---")
        st.subheader("🛒 Configura tu plan")
        plan = st.radio("Selecciona tu plan", ["BASICO", "PRO", "VIP"], horizontal=True)
        tipo = st.radio("Tipo de pago", ["Mensual", "Pago Único"], horizontal=True)
        forma_pago = st.radio("Forma de pago", ["Transferencia", "Efectivo"], horizontal=True)

        PAQUETES = {"BASICO": {"mensual": 149, "pago_unico": 1999},"PRO": {"mensual": 299, "pago_unico": 3499},"VIP": {"mensual": 399, "pago_unico": 4999}}
        tipo_key = "mensual" if tipo == "Mensual" else "pago_unico"
        total = PAQUETES[plan][tipo_key]

        col1, col2 = st.columns([2,1])
        with col1:
            st.metric("Total a pagar", f"${total:,.0f} MXN", delta=f"{tipo}")
        st.markdown("---")
        with st.form("registro"):
            st.subheader("📝 Datos de facturación")
            col1, col2 = st.columns(2)
            with col1:
                nombre = st.text_input("Nombre Completo *")
                email = st.text_input("Email *")
            with col2:
                tel = st.text_input("Teléfono/WhatsApp")
                empresa = st.text_input("Razón Social o Empresa")
            if st.form_submit_button("💳 Solicitar Alta", use_container_width=True, type="primary"):
                if nombre and email:
                    try:
                        r = requests.post(f"{SERVIDOR}/api/registrar_empresa", json={
                            "nombre": nombre,
                            "email": email,
                            "tel": tel,
                            "empresa": empresa,
                            "plan": plan,
                            "tipo_pago": tipo_key,
                            "forma_pago": forma_pago.lower()
                        }, timeout=10)
                        r.raise_for_status()
                        data = r.json()
                        if data['status'] == 'ok':
                            st.success(f"✅ {data['msg']}")
                            st.info("📧 Revisa tu correo. Te enviamos los datos de pago. Tu usuario se activará al confirmar la transferencia.")
                            st.balloons()
                        else:
                            st.error(f"❌ {data['msg']}")
                    except Exception as e:
                        st.error(f"❌ Error: {str(e)}")
                else:
                    st.error("Completa nombre y email obligatorios")

    with tab3:
        st.header("🏢 Sobre ContaXpert Pro")
        st.markdown("""
        ### 🎯 Nuestra Visión
        Ser la plataforma líder en México para la automatización fiscal y contable, empoderando a contadores y empresas para que dediquen su tiempo a lo que realmente importa: crecer su negocio.

        ### 🚀 Nuestro Objetivo
        Eliminar el trabajo manual repetitivo en la contabilidad mediante tecnología de punta, reduciendo errores humanos y acelerando los procesos fiscales hasta 10 veces.

        ### 💻 Lo Que Hacemos
        En **ContaXpert Pro** somos especialistas en automatización de procesos de negocio. Desarrollamos software inteligente para:
        - **Contabilidad:** Conversión masiva de XML a Excel, PDF y Word con análisis automático
        - **Facturación:** Integración con SAT y validación de CFDI en tiempo real
        - **Nómina:** Cálculo automático y timbrado masivo
        - **Inventarios:** Sincronización con sistemas ERP
        - **Reportes:** Dashboards ejecutivos con KPIs en tiempo real

        Nuestro equipo de programadores expertos en Python, APIs y automatización RPA transforma procesos manuales de horas en tareas de minutos.

        ### 🤝 ¿Por Qué Elegirnos?
        - ✅ **+500 despachos** confían en nosotros
        - ✅ **Soporte 24/7** por WhatsApp con respuesta en <30 min
        - ✅ **Actualizaciones gratuitas** cuando cambia el SAT
        - ✅ **Capacitación incluida** en todos los planes
        - ✅ **100% Hecho en México** 🇲🇽
        """)
    st.stop()

# ================= PANTALLA PRINCIPAL =================
datos = st.session_state.datos
if not datos:
    st.error("Error de sesión. Vuelve a iniciar sesión")
    if st.button("Volver al login"):
        st.session_state.usuario = None
        st.rerun()
    st.stop()

# Sidebar
st.sidebar.markdown(f'<div class="header-logo">{LOGO_SVG}</div>', unsafe_allow_html=True)
st.sidebar.success(f"👤 {datos['nombre']}")
st.sidebar.caption(f"Usuario: `{st.session_state.usuario}`")
st.sidebar.divider()
col1, col2 = st.sidebar.columns(2)
with col1:
    st.metric("Plan", datos['plan'])
with col2:
    if datos['ilimitado']:
        st.metric("XML", "♾️")
    else:
        st.metric("Disponibles", datos['disponibles'])
if not datos['ilimitado']:
    st.sidebar.progress(datos['usados'] / datos['limite'] if datos['limite'] > 0 else 0)
    st.sidebar.caption(f"Usados: {datos['usados']}/{datos['limite']} + {datos['extra']} extra")
if datos['vence'] == 'Vitalicio':
    st.sidebar.success("🎉 Pago Único - Sin vencimiento")
else:
    st.sidebar.caption(f"📅 Vence: {datos['vence']}")
st.sidebar.divider()
if st.sidebar.button("🚪 Cerrar Sesión", use_container_width=True):
    st.session_state.usuario = None
    st.session_state.datos = None
    st.session_state.modo_demo = False
    st.session_state.df_convertido = None
    st.session_state.descarga_contabilizada = False
    st.rerun()

# Main
st.markdown(f'<div class="header-logo">{LOGO_SVG}</div>', unsafe_allow_html=True)
st.title(f"Bienvenido {datos['nombre']} 👋")

tab1, tab2, tab3, tab4 = st.tabs(["📄 Convertir XML", "💬 Soporte", "🔄 Renovar", "🏢 Quiénes Somos"])

with tab1:
    st.header("Convertir XML a Excel/PDF/Word")

    with st.expander("⚙️ 1. Selecciona los campos que quieres en tu reporte", expanded=True):
        cols = st.columns(3)
        idx = 0
        for categoria, campos in CAMPOS_CFDI.items():
            with cols[idx % 3]:
                st.markdown(f"**{categoria}**")
                for campo in campos:
                    checked = st.checkbox(campo, key=f"campo_{campo}", value=campo in st.session_state.campos_seleccionados)
                    if checked and campo not in st.session_state.campos_seleccionados:
                        st.session_state.campos_seleccionados.append(campo)
                    elif not checked and campo in st.session_state.campos_seleccionados:
                        st.session_state.campos_seleccionados.remove(campo)
            idx += 1

    if not st.session_state.campos_seleccionados:
        st.warning("Selecciona al menos 1 campo para continuar")
        st.stop()

    st.markdown("---")
    st.subheader("📁 2. Carga tus archivos XML")

    if not datos['ilimitado'] and datos['disponibles'] <= 0:
        st.error("❌ Sin XML disponibles este mes")
        st.info("💡 Contacta a soporte para renovar tu plan")
    else:
        archivos = st.file_uploader("Arrastra tus archivos XML aquí", accept_multiple_files=True, type=['xml'], key="uploader_key")

        if archivos:
            st.write(f"📁 Archivos seleccionados: {len(archivos)}")
            if not datos['ilimitado'] and len(archivos) > datos['disponibles']:
                st.error(f"❌ Solo quedan {datos['disponibles']} XML disponibles")
            else:
                if st.button("⚡ Convertir Ahora", type="primary", use_container_width=True):
                    with st.spinner('Procesando XMLs... Esto puede tardar unos segundos'):
                        try:
                            df = procesar_xmls_detalle(archivos, st.session_state.campos_seleccionados)
                            st.session_state.df_convertido = df
                            st.session_state.archivos_temp = archivos
                            st.session_state.descarga_contabilizada = False
                            st.success(f"✅ {len(archivos)} archivos convertidos - {len(df)} filas generadas")
                            st.rerun()
                        except Exception as e:
                            st.error(f"❌ Error al procesar: {str(e)}")

    if st.session_state.df_convertido is not None:
        df = st.session_state.df_convertido
        st.markdown("---")
        st.subheader("📊 3. Resultados")

        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Total Facturas", df['UUID'].nunique() if 'UUID' in df.columns else len(df))
        with col2:
            if 'Total' in df.columns:
                st.metric("Monto Total", f"${df['Total'].sum():,.2f}")
        with col3:
            if 'SubTotal' in df.columns:
                st.metric("Subtotal", f"${df['SubTotal'].sum():,.2f}")
        with col4:
            if 'IVA_Trasladado' in df.columns:
                st.metric("IVA Total", f"${df['IVA_Trasladado'].sum():,.2f}")

        if 'Total' in df.columns and 'Nombre_Emisor' in df.columns:
            col1, col2 = st.columns(2)
            with col1:
                fig1 = px.bar(df.groupby('Nombre_Emisor')['Total'].sum().reset_index().head(10), x='Nombre_Emisor', y='Total', title='Top 10 Emisores por Monto', color_discrete_sequence=['#667eea'])
                st.plotly_chart(fig1, use_container_width=True)
            with col2:
                if 'Fecha' in df.columns:
                    df_temp = df.copy()
                    df_temp['Mes'] = pd.to_datetime(df_temp['Fecha'], errors='coerce').dt.to_period('M').astype(str)
                    fig2 = px.line(df_temp.groupby('Mes')['Total'].sum().reset_index(), x='Mes', y='Total', title='Tendencia Mensual', color_discrete_sequence=['#764ba2'])
                    st.plotly_chart(fig2, use_container_width=True)

        st.subheader("👁️ Vista Previa")
        st.dataframe(df.head(100), use_container_width=True, height=300)

        st.subheader("💾 4. Descargar Reporte")
        col1, col2 = st.columns([1,1])
        with col1:
            excel_bytes = generar_excel(df)
            if st.download_button("📥 Descargar Excel", excel_bytes, file_name=f"ContaXpert_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True):
                # CAMBIO QUIRÚRGICO: Descontar solo la primera vez que descarga
                if not st.session_state.descarga_contabilizada:
                    try:
                        cantidad = len(st.session_state.archivos_temp)
                        r = requests.post(f"{SERVIDOR}/api/usar", json={"usuario": st.session_state.usuario, "cantidad": cantidad}, timeout=15)
                        r.raise_for_status()
                        if r.json()['ok']:
                            st.session_state.datos['disponibles'] -= cantidad
                            st.session_state.datos['usados'] += cantidad
                            st.session_state.descarga_contabilizada = True
                            st.success(f"✅ Se descontaron {cantidad} XML del plan")
                            # CAMBIO QUIRÚRGICO: Limpiar archivos para poder subir más
                            st.session_state.archivos_temp = []
                            st.session_state.df_convertido = None
                            st.rerun()
                        else:
                            st.error(f"❌ {r.json()['msg']}")
                    except Exception as e:
                        st.error(f"❌ Error al descontar: {str(e)}")
        with col2:
            st.info("ℹ️ El descuento se aplica solo la primera vez que descargas. Puedes descargar el archivo las veces que quieras.")

with tab2:
    st.header("Soporte Técnico Profesional")
    st.caption(f"Mensaje de: {datos['nombre']} | Usuario: {st.session_state.usuario} | Plan: {datos['plan']}")

    col1, col2 = st.columns(2)
    with col1:
        asunto = st.selectbox("Tipo de problema", [
            "Error al convertir XML",
            "Problema con pagos",
            "Duda de facturación",
            "Error de login",
            "Otro"
        ])
    with col2:
        urgencia = st.radio("Urgencia", ["Normal", "Urgente"], horizontal=True)

    mensaje = st.text_area("Describe tu duda o problema con detalle", height=150, placeholder="Ej: No puedo descargar el PDF, me marca error al cargar 50 XMLs...")

    if st.button("📤 Enviar mensaje", use_container_width=True, type="primary"):
        if mensaje:
            payload = {
                "usuario": st.session_state.usuario,
                "nombre": datos['nombre'],
                "email": datos.get('email', 'No registrado'),
                "plan": datos['plan'],
                "asunto": asunto,
                "urgencia": urgencia,
                "mensaje": mensaje
            }
            try:
                r = requests.post(f"{SERVIDOR}/api/enviar_mensaje", json=payload, timeout=10)
                r.raise_for_status()
                res = r.json()
                if res['status'] == 'ok':
                    st.success("✅ Mensaje enviado exitosamente")
                    st.info("📧 Recibirás una copia en tu correo. Nuestro equipo te contactará en menos de 30 min si es urgente.")
                    st.balloons()
                else:
                    st.error(f"❌ {res['msg']}")
            except Exception as e:
                st.error(f"❌ Error al enviar: {str(e)}")
        else:
            st.error("Por favor escribe tu mensaje")

with tab3:
    st.header("Renovar Suscripción")
    st.info("Para renovar contacta a soporte: wa.me/526331124596")
    st.write("Te enviaremos los datos de pago por transferencia o efectivo.")

with tab4:
    st.header("🏢 Sobre ContaXpert Pro")
    st.markdown("""
    ### 🎯 Nuestra Visión
    Ser la plataforma líder en México para la automatización fiscal y contable, empoderando a contadores y empresas para que dediquen su tiempo a lo que realmente importa: **crecer su negocio**.

    ### 🚀 Nuestro Objetivo
    Eliminar el trabajo manual repetitivo en la contabilidad mediante tecnología de punta, reduciendo errores humanos y acelerando los procesos fiscales hasta **10 veces**.

    ### 💻 Lo Que Hacemos
    En **ContaXpert Pro** somos especialistas en automatización de procesos de negocio. Desarrollamos software inteligente para:

    - **📊 Contabilidad:** Conversión masiva de XML a Excel, PDF y Word con análisis automático
    - **🧾 Facturación:** Integración con SAT y validación de CFDI en tiempo real
    - **👥 Nómina:** Cálculo automático y timbrado masivo
    - **📦 Inventarios:** Sincronización con sistemas ERP
    - **📈 Reportes:** Dashboards ejecutivos con KPIs en tiempo real

    Nuestro equipo de programadores expertos en Python, APIs y automatización RPA transforma procesos manuales de horas en tareas de minutos.

    ### 🤝 ¿Por Qué Elegirnos?
    - ✅ **+500 despachos** confían en nosotros
    - ✅ **Soporte 24/7** por WhatsApp con respuesta en <30 min
    - ✅ **Actualizaciones gratuitas** cuando cambia el SAT
    - ✅ **Capacitación incluida** en todos los planes
    - ✅ **100% Hecho en México** 🇲🇽
    """)
